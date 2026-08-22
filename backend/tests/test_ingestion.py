import io
import fitz
import pytest
from unittest.mock import MagicMock, patch
from app.config.settings import settings
from app.core.ingestion.s3_loader import S3Loader
from app.core.ingestion.parser import parse_bytes
from app.core.ingestion.pipeline import _AdaptiveLimiter, _CpuGovernor


_HEADING_FONT = 24
_BODY_FONT = 11
_CELL_FONT = 10


def _make_pdf(text: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _draw_table(page, cells: list[list[str]], origin=(50, 140), cell=(120, 24)) -> None:
    """Draw a grid + cell text so pymupdf4llm detects a real table."""
    x0, y0 = origin
    cw, rh = cell
    rows, cols = len(cells), len(cells[0])
    for r in range(rows + 1):
        page.draw_line((x0, y0 + r * rh), (x0 + cols * cw, y0 + r * rh))
    for c in range(cols + 1):
        page.draw_line((x0 + c * cw, y0), (x0 + c * cw, y0 + rows * rh))
    for r in range(rows):
        for c in range(cols):
            page.insert_text((x0 + c * cw + 5, y0 + r * rh + 16), cells[r][c], fontsize=_CELL_FONT)


def _make_structured_pdf(heading: str, body: str, cells: list[list[str]]) -> bytes:
    """A single-page PDF with a large-font heading, body text, and a drawn table."""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 60), heading, fontsize=_HEADING_FONT)
    page.insert_text((50, 100), body, fontsize=_BODY_FONT)
    _draw_table(page, cells)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_s3_loader_local_fallback(tmp_path, monkeypatch):
    # Isolate from a live S3_BUCKET_NAME in the local .env so the loader takes
    # the local-filesystem fallback path instead of hitting real S3.
    monkeypatch.setattr(settings, "S3_BUCKET_NAME", None)
    monkeypatch.setattr(settings, "S3_BUCKET_NAMES", None)
    (tmp_path / "a.pdf").write_bytes(b"%PDF")
    loader = S3Loader(local_root=str(tmp_path))
    keys = loader.list_keys()
    assert "a.pdf" in keys


def test_s3_loader_download_local(tmp_path):
    (tmp_path / "doc.pdf").write_bytes(b"content")
    loader = S3Loader(local_root=str(tmp_path))
    assert loader.download("doc.pdf") == b"content"


def test_parse_pdf_returns_chunks():
    data = _make_pdf("Section 302 IPC defines murder and its punishment.")
    chunks = parse_bytes(data, "ipc.pdf")
    assert len(chunks) >= 1
    assert all("text" in c and "source" in c and "page" in c for c in chunks)
    assert chunks[0]["source"] == "ipc.pdf"
    assert "Section 302" in " ".join(c["text"] for c in chunks)


def test_parse_blank_pdf_returns_empty():
    doc = fitz.open()
    doc.new_page()
    buf = io.BytesIO()
    doc.save(buf)
    assert parse_bytes(buf.getvalue(), "blank.pdf") == []


def test_parse_txt_splits():
    text = "word " * 300
    chunks = parse_bytes(text.encode(), "doc.txt")
    assert len(chunks) >= 2


def test_parse_csv_renders_as_markdown_table():
    csv_bytes = b"Section,Description\n302,Punishment for murder\n304,Culpable homicide\n"
    chunks = parse_bytes(csv_bytes, "sections.csv")
    assert len(chunks) >= 1
    combined = " ".join(c["text"] for c in chunks)
    assert "Punishment for murder" in combined
    assert "| Section | Description |" in combined


def test_parse_csv_empty_returns_no_chunks():
    assert parse_bytes(b"", "empty.csv") == []


def test_parse_docx_extracts_paragraphs_and_headings():
    from io import BytesIO
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Section 302", level=1)
    doc.add_paragraph("Whoever commits murder shall be punished with death.")
    buf = BytesIO()
    doc.save(buf)

    chunks = parse_bytes(buf.getvalue(), "judgment.docx")
    combined = " ".join(c["text"] for c in chunks)
    assert "Section 302" in combined
    assert "Whoever commits murder" in combined


def test_parse_docx_extracts_tables():
    from io import BytesIO
    from docx import Document as DocxDocument

    doc = DocxDocument()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Statute"
    table.rows[0].cells[1].text = "Punishment"
    table.rows[1].cells[0].text = "Section 302"
    table.rows[1].cells[1].text = "Death or life imprisonment"
    buf = BytesIO()
    doc.save(buf)

    chunks = parse_bytes(buf.getvalue(), "schedule.docx")
    combined = " ".join(c["text"] for c in chunks)
    assert "Death or life imprisonment" in combined


def test_parse_docx_empty_returns_no_chunks():
    from io import BytesIO
    from docx import Document as DocxDocument

    buf = BytesIO()
    DocxDocument().save(buf)
    assert parse_bytes(buf.getvalue(), "empty.docx") == []


def test_parse_image_ocrs_when_text_found(monkeypatch):
    """OCR stack (pytesseract) is mocked — this test verifies parse_bytes'
    dispatch + chunking for images, not real OCR accuracy, since the
    tesseract binary isn't installed on every dev/CI machine."""
    import app.core.ingestion.parser as parser_mod

    fake_pytesseract = MagicMock()
    fake_pytesseract.image_to_string.return_value = "Section 302 defines murder. " * 10
    monkeypatch.setitem(__import__("sys").modules, "pytesseract", fake_pytesseract)

    from PIL import Image
    img = Image.new("RGB", (100, 100), color="white")
    from io import BytesIO
    buf = BytesIO()
    img.save(buf, format="PNG")

    chunks = parser_mod.parse_bytes(buf.getvalue(), "scan.png")
    assert len(chunks) >= 1
    assert "Section 302" in " ".join(c["text"] for c in chunks)


def test_parse_image_below_min_chars_returns_empty(monkeypatch):
    import app.core.ingestion.parser as parser_mod

    fake_pytesseract = MagicMock()
    fake_pytesseract.image_to_string.return_value = "x"
    monkeypatch.setitem(__import__("sys").modules, "pytesseract", fake_pytesseract)

    from PIL import Image
    img = Image.new("RGB", (10, 10), color="white")
    from io import BytesIO
    buf = BytesIO()
    img.save(buf, format="PNG")

    assert parser_mod.parse_bytes(buf.getvalue(), "blank.png") == []


def test_parse_image_ocr_disabled_returns_empty(monkeypatch):
    monkeypatch.setattr(settings, "OCR_ENABLED", False)
    from PIL import Image
    img = Image.new("RGB", (100, 100), color="white")
    from io import BytesIO
    buf = BytesIO()
    img.save(buf, format="PNG")

    assert parse_bytes(buf.getvalue(), "scan.png") == []


def test_parse_pdf_preserves_markdown_structure():
    """Headings (#) and table syntax (|) survive extraction into chunk text."""
    data = _make_structured_pdf(
        heading="Indian Penal Code",
        body="Section 302 defines the punishment for murder.",
        cells=[["Section", "Punishment"], ["302", "Life / Death"], ["304", "Imprisonment"]],
    )
    chunks = parse_bytes(data, "ipc.pdf")
    combined = "\n".join(c["text"] for c in chunks)
    assert "# Indian Penal Code" in combined
    assert "|Section|Punishment|" in combined
    assert "|---" in combined  # markdown table separator row
    assert "302" in combined and "Life / Death" in combined


def test_parse_pdf_structure_metadata_correct():
    """Every structured chunk carries the right source and 0-based page index."""
    data = _make_structured_pdf(
        heading="Contract Law",
        body="An offer must be accepted to form a binding agreement.",
        cells=[["Element", "Required"], ["Offer", "Yes"], ["Acceptance", "Yes"]],
    )
    chunks = parse_bytes(data, "contracts.pdf")
    assert chunks, "expected at least one chunk from a structured PDF"
    assert all(c["source"] == "contracts.pdf" for c in chunks)
    assert all(c["page"] == 0 for c in chunks)


@pytest.mark.asyncio
async def test_adaptive_limiter_grows_on_success():
    limiter = _AdaptiveLimiter(initial=2, minimum=2, maximum=8)
    await limiter.acquire()
    await limiter.release(ok=True)
    assert limiter.limit == 3
    await limiter.acquire()
    await limiter.release(ok=True)
    assert limiter.limit == 4


@pytest.mark.asyncio
async def test_adaptive_limiter_halves_on_failure_and_floors_at_minimum():
    limiter = _AdaptiveLimiter(initial=8, minimum=2, maximum=24)
    await limiter.acquire()
    await limiter.release(ok=False)
    assert limiter.limit == 4
    await limiter.acquire()
    await limiter.release(ok=False)
    assert limiter.limit == 2  # floored at minimum, not 1
    await limiter.acquire()
    await limiter.release(ok=False)
    assert limiter.limit == 2


@pytest.mark.asyncio
async def test_adaptive_limiter_caps_at_maximum():
    limiter = _AdaptiveLimiter(initial=8, minimum=2, maximum=8)
    await limiter.acquire()
    await limiter.release(ok=True)
    assert limiter.limit == 8  # already at ceiling, success doesn't overshoot


def test_cpu_governor_degrades_above_budget_and_floors_at_minimum():
    gov = _CpuGovernor(budget_percent=50.0, min_ceiling=2, max_ceiling=8, sample_interval=2.0)
    gov.ceiling = 8
    gov._step(90.0)
    assert gov.ceiling == 7
    for _ in range(10):
        gov._step(90.0)
    assert gov.ceiling == 2  # never drops below minimum — pipeline stays alive


def test_cpu_governor_recovers_when_load_drops():
    gov = _CpuGovernor(budget_percent=50.0, min_ceiling=2, max_ceiling=8, sample_interval=2.0)
    gov.ceiling = 2
    gov._step(10.0)
    assert gov.ceiling == 3


@pytest.mark.asyncio
async def test_limiter_is_clamped_by_governor_ceiling():
    gov = _CpuGovernor(budget_percent=50.0, min_ceiling=2, max_ceiling=8, sample_interval=2.0)
    gov.ceiling = 3
    limiter = _AdaptiveLimiter(initial=8, minimum=2, maximum=8, governor=gov)
    # AIMD's internal limit is 8, but the governor's CPU-budget ceiling (3) wins.
    assert limiter.limit == 3


@pytest.mark.asyncio
async def test_list_pdf_keys_times_out_without_hanging(monkeypatch):
    """A listing call that never returns (a stuck botocore call — DNS, a
    dead connection, a retry loop) must not hang the caller forever. This
    replaced a forked-subprocess design that, in practice, deadlocked on
    every call when forked from a live multi-threaded async server (a lock
    held by another thread at fork time is inherited "locked forever" in
    the child). asyncio.wait_for over a plain to_thread call has no such
    fork risk and still bounds the wait."""
    import threading
    from app.core.ingestion import pipeline

    release = threading.Event()

    def _hangs_forever(prefix_filter):
        release.wait()  # never set within the test — simulates a stuck call
        return []

    monkeypatch.setattr(pipeline, "_list_pdf_keys", _hangs_forever)

    with pytest.raises(TimeoutError):
        await pipeline._list_pdf_keys_with_timeout("", timeout=0.05)

    release.set()  # let the orphaned thread-pool worker finish so it doesn't leak


def test_parse_pdf_page_metadata_increments_across_pages():
    """Chunks report the page they came from across a multi-page document."""
    doc = fitz.open()
    doc.new_page().insert_text((50, 60), "Volume One", fontsize=_HEADING_FONT)
    doc.new_page().insert_text((50, 60), "Volume Two", fontsize=_HEADING_FONT)
    buf = io.BytesIO()
    doc.save(buf)
    chunks = parse_bytes(buf.getvalue(), "multi.pdf")
    pages = {c["page"] for c in chunks}
    assert pages == {0, 1}


# ---------------------------------------------------------------------------
# Postgres mirror of the corpus — SpiceAI has no Qdrant connector, so this
# table is what its SQL/NQL/semantic search against the main corpus
# actually queries. See app.models.legal_chunk.LegalChunk.
# ---------------------------------------------------------------------------

def test_chunk_content_id_matches_qdrant_quickwit_formula():
    """The three stores must agree on a chunk's identity — same
    uuid5(source|page|text) — so a re-ingest overwrites the same row/point
    everywhere instead of drifting into three different ids."""
    import uuid
    from app.core.ingestion.pipeline import _chunk_content_id

    expected = str(uuid.uuid5(uuid.NAMESPACE_URL, "ipc.pdf|3|Section 302 defines murder."))
    assert _chunk_content_id("ipc.pdf", 3, "Section 302 defines murder.") == expected


@pytest.mark.asyncio
async def test_upsert_legal_chunks_pg_is_noop_for_empty_chunks():
    from unittest.mock import AsyncMock
    from app.core.ingestion import pipeline

    with patch.object(pipeline, "get_sessionmaker") as mock_sessionmaker:
        await pipeline._upsert_legal_chunks_pg([])

    mock_sessionmaker.assert_not_called()


@pytest.mark.asyncio
async def test_upsert_legal_chunks_pg_executes_upsert_and_commits():
    from unittest.mock import AsyncMock
    from app.core.ingestion import pipeline

    mock_db = MagicMock()
    mock_db.execute = AsyncMock()
    mock_db.commit = AsyncMock()

    class _FakeSessionCtx:
        async def __aenter__(self):
            return mock_db
        async def __aexit__(self, *exc):
            return False

    mock_sessionmaker = MagicMock(return_value=_FakeSessionCtx())

    with patch.object(pipeline, "get_sessionmaker", return_value=mock_sessionmaker):
        await pipeline._upsert_legal_chunks_pg(
            [{"source": "ipc.pdf", "page": 3, "text": "Section 302 defines murder."}]
        )

    mock_db.execute.assert_called_once()
    mock_db.commit.assert_called_once()
    stmt = mock_db.execute.call_args[0][0]
    # The compiled statement targets the legal_chunks table, upserting the
    # same content-addressed id Qdrant/Quickwit computed for this chunk.
    assert "legal_chunks" in str(stmt)
