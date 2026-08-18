from __future__ import annotations
import asyncio
import json
import logging
import time
import uuid
from datetime import datetime
from io import BytesIO
from typing import AsyncGenerator, Optional, List, Dict, Any

from fastapi import APIRouter, Request, HTTPException, Body
from fastapi.responses import StreamingResponse, Response
from pydantic import BaseModel
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_core.messages import HumanMessage

from app.api.schemas import (
    ChatAnalyzeResponse,
    ChatImproveResponse,
    ChatDownloadPdfRequest,
    ChatRequest,
    ChatResponse,
    CitationOut,
    DevilsAdvocateRequest,
    DevilsAdvocateResponse,
    LlmStatusOut,
    SourceChunkOut,
    VerificationOut,
    QueryAnalysisOut,
)
from app.config.settings import settings
from app.core.graph.workflow import (
    get_workflow, get_streaming_workflow, get_enhanced_workflow, 
    emit_progress, PROGRESS_STEPS
)
from app.core.graph.state import JuryAIState
from app.core.graph.verifier import verify_answer, _build_evidence_text
from app.core.graph.gate import gate_answer
from app.core.retrieval.citation import derive_citations, verified_content_hashes, cited_indices
from app.core.llm.provider import get_llm
from app.core.cache import answer_cache_get, answer_cache_set
from app.core.memory import load_history, append_turn
from app.core.ratelimit import check_rate_limit
from app.core.audit import build_turn_record, schedule_record_turn
from app.core.run_store import save_run, load_run, list_runs

router = APIRouter()

logger = logging.getLogger("juryai.chat")


@router.get("/llm/status", response_model=LlmStatusOut)
async def llm_status():
    """The provider actually in effect — settings.llm_provider_chain[0], the
    one every chat request tries first — so Settings can show the real
    connected gateway instead of a hardcoded description."""
    chain = settings.llm_provider_chain
    if not chain:
        return LlmStatusOut(provider="none", model="", base_url=None, configured=False)
    active = chain[0]
    return LlmStatusOut(
        provider=active["provider_name"],
        model=active["model"],
        base_url=active.get("base_url"),
        configured=True,
        web_search_provider=settings.web_search_provider,
    )


# ---------------------------------------------------------------------------
# PDF Generation Helpers
# ---------------------------------------------------------------------------

class _PDF(FPDF):
    _FONT_DIR = "/usr/share/fonts/truetype/liberation"
    
    def __init__(self):
        super().__init__()
        self.add_font("LiberationSans", "", f"{self._FONT_DIR}/LiberationSans-Regular.ttf", uni=True)
        self.add_font("LiberationSans", "B", f"{self._FONT_DIR}/LiberationSans-Bold.ttf", uni=True)
        self.add_font("LiberationSans", "I", f"{self._FONT_DIR}/LiberationSans-Italic.ttf", uni=True)
        self.add_font("LiberationSans", "BI", f"{self._FONT_DIR}/LiberationSans-BoldItalic.ttf", uni=True)
    
    def header(self):
        self.set_font("LiberationSans", "B", 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 6, "JuryAI Legal Opinion", align="C", new_x="LMARGIN", new_y="NEXT")
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("LiberationSans", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def _generate_pdf(
    question: str,
    answer: str,
    citations: List[CitationOut],
    source_chunks: List[SourceChunkOut],
    verification: Optional[VerificationOut],
    include_citations: bool = True,
) -> bytes:
    pdf = _PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    pdf.set_font("LiberationSans", "B", 16)
    pdf.set_text_color(30, 30, 30)
    pdf.multi_cell(0, 8, "Legal Opinion", align="C")
    pdf.ln(4)

    pdf.set_font("LiberationSans", "B", 11)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(0, 7, "Question:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("LiberationSans", "", 11)
    pdf.set_text_color(30, 30, 30)
    pdf.multi_cell(0, 6, question)
    pdf.ln(6)

    pdf.set_font("LiberationSans", "B", 11)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(0, 7, "Answer:", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("LiberationSans", "", 11)
    pdf.set_text_color(30, 30, 30)

    if include_citations and citations:
        pdf.multi_cell(0, 6, answer)
    else:
        import re
        clean_answer = re.sub(r"\s*\[\d+(?:,\s*\d+)*\]", "", answer)
        pdf.multi_cell(0, 6, clean_answer)
    pdf.ln(6)

    if verification:
        pdf.set_font("LiberationSans", "B", 11)
        pdf.set_text_color(50, 50, 50)
        pdf.cell(0, 7, "Verification:", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("LiberationSans", "", 10)
        verdict_color = {
            "grounded": (0, 128, 0),
            "partially_grounded": (200, 150, 0),
            "unsupported": (200, 0, 0),
        }.get(verification.verdict, (128, 128, 128))
        pdf.set_text_color(*verdict_color)
        pdf.cell(0, 6, f"Verdict: {verification.verdict.replace('_', ' ').title()}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 6, f"Groundedness: {verification.groundedness_score:.0%}", new_x="LMARGIN", new_y="NEXT")
        if verification.summary:
            pdf.multi_cell(0, 6, f"Summary: {verification.summary}")
        pdf.ln(4)

    if include_citations and citations:
        pdf.set_font("LiberationSans", "B", 11)
        pdf.set_text_color(50, 50, 50)
        pdf.cell(0, 7, "Citations:", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for cit in citations:
            pdf.set_font("LiberationSans", "B", 10)
            pdf.set_text_color(30, 30, 30)
            marker = f"[{cit.index}]"
            if cit.verified:
                marker += " \u2713"
            pdf.cell(0, 6, marker, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("LiberationSans", "", 9)
            pdf.set_text_color(60, 60, 60)
            pdf.multi_cell(0, 5, f'"{cit.quote[:200]}{"..." if len(cit.quote) > 200 else ""}"')
            pdf.set_font("LiberationSans", "I", 9)
            pdf.cell(0, 5, f"Source: {cit.source}", new_x="LMARGIN", new_y="NEXT")
            if cit.page:
                pdf.cell(0, 5, f"Page: {cit.page}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

    if include_citations and source_chunks:
        pdf.add_page()
        pdf.set_font("LiberationSans", "B", 11)
        pdf.set_text_color(50, 50, 50)
        pdf.cell(0, 7, "Sources:", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for sc in source_chunks:
            pdf.set_font("LiberationSans", "B", 10)
            pdf.set_text_color(30, 30, 30)
            domain_label = "Internal" if sc.domain == "internal" else "Web"
            verified_marker = " \u2713" if sc.verified else ""
            cited_marker = " [cited]" if sc.cited else ""
            pdf.cell(0, 6, f"[{sc.index}] {sc.source} ({domain_label}){verified_marker}{cited_marker}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("LiberationSans", "", 9)
            pdf.set_text_color(60, 60, 60)
            pdf.multi_cell(0, 5, sc.text[:300] + ("..." if len(sc.text) > 300 else ""))
            if sc.url:
                pdf.set_font("LiberationSans", "I", 9)
                pdf.set_text_color(0, 0, 200)
                pdf.cell(0, 5, f"URL: {sc.url}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

    pdf.ln(10)
    pdf.set_font("LiberationSans", "I", 8)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 5, f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 5, "JuryAI - Legal Research Assistant", align="C")

    return bytes(pdf.output())


_VERDICT_COLOR = {
    "grounded": RGBColor(0, 128, 0),
    "partially_grounded": RGBColor(200, 150, 0),
    "unsupported": RGBColor(200, 0, 0),
}


def _generate_docx(
    question: str,
    answer: str,
    citations: List[CitationOut],
    source_chunks: List[SourceChunkOut],
    verification: Optional[VerificationOut],
    include_citations: bool = True,
) -> bytes:
    doc = Document()

    title = doc.add_heading("Legal Opinion", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Question:", level=2)
    doc.add_paragraph(question)

    doc.add_heading("Answer:", level=2)
    if include_citations and citations:
        doc.add_paragraph(answer)
    else:
        import re
        doc.add_paragraph(re.sub(r"\s*\[\d+(?:,\s*\d+)*\]", "", answer))

    if verification:
        doc.add_heading("Verification:", level=2)
        p = doc.add_paragraph()
        run = p.add_run(f"Verdict: {verification.verdict.replace('_', ' ').title()}")
        run.font.color.rgb = _VERDICT_COLOR.get(verification.verdict, RGBColor(128, 128, 128))
        doc.add_paragraph(f"Groundedness: {verification.groundedness_score:.0%}")
        if verification.summary:
            doc.add_paragraph(f"Summary: {verification.summary}")

    if include_citations and citations:
        doc.add_heading("Citations:", level=2)
        for cit in citations:
            marker = f"[{cit.index}]" + (" ✓" if cit.verified else "")
            p = doc.add_paragraph()
            p.add_run(marker).bold = True
            doc.add_paragraph(f'"{cit.quote[:200]}{"..." if len(cit.quote) > 200 else ""}"')
            source_line = f"Source: {cit.source}" + (f" — Page: {cit.page}" if cit.page else "")
            doc.add_paragraph(source_line).runs[0].italic = True

    if include_citations and source_chunks:
        doc.add_page_break()
        doc.add_heading("Sources:", level=2)
        for sc in source_chunks:
            domain_label = "Internal" if sc.domain == "internal" else "Web"
            verified_marker = " ✓" if sc.verified else ""
            cited_marker = " [cited]" if sc.cited else ""
            p = doc.add_paragraph()
            p.add_run(f"[{sc.index}] {sc.source} ({domain_label}){verified_marker}{cited_marker}").bold = True
            doc.add_paragraph(sc.text[:300] + ("..." if len(sc.text) > 300 else ""))
            if sc.url:
                doc.add_paragraph(f"URL: {sc.url}").runs[0].italic = True

    footer = doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _sse(event: str, data: dict) -> str:
    """Format a single SSE message."""
    return f"event: {event}\ndata: {json.dumps(data)}\n\n"


def _client_id(http_request: Request) -> str:
    """Best-effort client identifier for rate limiting.

    ponytail: trusts X-Forwarded-For's first hop — fine behind our own reverse
    proxy; tighten (trusted-proxy allowlist) if exposed directly to clients.
    """
    xff = http_request.headers.get("x-forwarded-for")
    if xff:
        return xff.split(",")[0].strip()
    return http_request.client.host if http_request.client else "unknown"


def _build_source_chunks(
    result: dict, verified_hashes: set[str] = None, cited_idx: set[int] = frozenset(), citations: list[dict] = None
) -> tuple[list[SourceChunkOut], list[dict]]:
    """
    Legacy function for backward compatibility with tests.

    Builds source chunks from legal_chunks (pre-merge evidence) only.
    Always sets domain="internal" since pre-merge evidence doesn't have domain tags.
    Returns (chunks, citations_raw) for compatibility with test assertions.
    
    Accepts either a dict with "legal_chunks" key or a list directly.
    """
    if verified_hashes is None:
        verified_hashes = set()
    
    # Handle both dict with "legal_chunks" key and direct list
    if isinstance(result, dict):
        evidence = result.get("legal_chunks", [])
        citations = result.get("citations", [])
    else:
        evidence = result
        citations = citations or []
    
    citation_quotes = {}
    if citations:
        for cit in citations:
            if cit.get("index") and cit.get("quote"):
                citation_quotes[cit["index"]] = cit["quote"]

    chunks: list[SourceChunkOut] = []
    for i, item in enumerate(evidence, 1):
        text = item.get("text") or item.get("content") or ""
        if not text:
            continue
        source = item.get("source")
        if source is None:
            source = item.get("title") or item.get("url") or ""
        # Legacy: pre-merge evidence is always internal
        domain = "internal"
        chunks.append(
            SourceChunkOut(
                text=text,
                source=source,
                page=item.get("page", 0),
                score=round(item.get("final_score", item.get("score", 0.0)), 4),
                verified=bool(item.get("content_hash")) and item.get("content_hash") in verified_hashes,
                cited=i in cited_idx if cited_idx else False,
                domain=domain,
                url=None,  # internal chunks never have URLs
                index=i,
                citation_quote=citation_quotes.get(i, ""),
                preview=text[:200] + ("..." if len(text) > 200 else ""),
                doc_id=item.get("doc_id"),
                chunk_id=item.get("chunk_id"),
            )
        )
    return chunks, citations or []


def _source_chunks_from_evidence(
    evidence: list[dict], verified_hashes: set[str], cited_idx: set[int] = frozenset(), citations: list[dict] = None
) -> list[SourceChunkOut]:
    citation_quotes = {}
    if citations:
        for cit in citations:
            if cit.get("index") and cit.get("quote"):
                citation_quotes[cit["index"]] = cit["quote"]

    chunks: list[SourceChunkOut] = []
    for i, item in enumerate(evidence, 1):
        text = item.get("text") or item.get("content") or ""
        if not text:
            continue
        source = item.get("source")
        if source is None:
            source = item.get("title") or item.get("url") or ""
        domain = item.get("domain", "internal")
        chunks.append(
            SourceChunkOut(
                text=text,
                source=source,
                page=item.get("page", 0),
                score=round(item.get("final_score", item.get("score", 0.0)), 4),
                verified=bool(item.get("content_hash")) and item.get("content_hash") in verified_hashes,
                cited=i in cited_idx,
                domain=domain,
                url=item.get("url") if domain == "web" else None,
                index=i,
                citation_quote=citation_quotes.get(i, ""),
                preview=text[:200] + ("..." if len(text) > 200 else ""),
                doc_id=item.get("doc_id"),
                chunk_id=item.get("chunk_id"),
            )
        )
    return chunks


async def _run_enhanced_chat(request: ChatRequest, http_request: Request) -> ChatResponse:
    """Run enhanced chat with run_id, progress tracking, structured output."""
    await _enforce_rate_limit(http_request)
    _validate_mode(request)
    
    conversation_id = request.conversation_id or str(uuid.uuid4())
    history = await load_history(conversation_id)
    
    # Use enhanced workflow
    graph = get_enhanced_workflow()
    
    state: JuryAIState = {
        "question": request.question,
        "history": history,
        "intent": "",
        "legal_chunks": [],
        "web_results": [],
        "citations": [],
        "answer": "",
        "error": None,
        "use_web_search": request.use_web_search,
        "mode": request.mode,
        "session_id": request.session_id,
        "output_format": request.output_format or "CREAC",
        "streaming": False,
        "source_filter": request.sources,
        "as_of_date": request.as_of_date,
    }
    
    result = await graph.ainvoke(state)
    
    answer = result.get("answer", "")
    verification = dict(result.get("verification") or {})
    citations_raw = result.get("citations", [])
    merged = result.get("merged_evidence") or []
    evidence = list(merged) if merged else list(result.get("legal_chunks") or [])
    
    # Verifier gate
    if settings.VERIFIER_GATE_ENABLED and evidence:
        gated = await gate_answer(
            request.question, answer, evidence, _build_evidence_text(evidence), verification
        )
        verification = dict(gated.verification)
        verification["blocked"] = gated.blocked
        verification["regenerated"] = gated.regenerated
        if not gated.blocked and gated.answer != answer:
            answer = gated.answer
            citations_raw = derive_citations(gated.verification, evidence, answer)
    
    citations = [CitationOut(**c) for c in citations_raw]
    verified_hashes = verified_content_hashes(citations_raw)
    source_chunks = _source_chunks_from_evidence(evidence, verified_hashes, cited_indices(citations_raw), citations_raw)
    verification_out = VerificationOut(**verification)
    
    # Save to history
    await append_turn(conversation_id, request.question, answer)
    schedule_record_turn(
        build_turn_record(
            conversation_id=conversation_id,
            question=request.question,
            answer=answer,
            citations=citations_raw,
            verification=verification,
            evidence=evidence,
            model_provider=result.get("model_provider", ""),
            model_name=result.get("model_name", ""),
            client_id=http_request.client.host if http_request.client else "unknown",
            streaming=False,
            cache_hit=False,
        )
    )
    
    response = ChatResponse(
        answer=answer,
        citations=citations,
        source_chunks=source_chunks,
        verification=verification_out,
        intent=result.get("intent", ""),
        sources_used=len(result.get("sources_used", {})) if isinstance(result.get("sources_used"), dict) else (result.get("sources_used") or 0),
        conversation_id=conversation_id,
        run_id=result.get("run_id"),
        output_format=result.get("output_format", "CREAC"),
        query_analysis=result.get("query_analysis"),
        reasoning_steps=result.get("reasoning_steps", []),
    )

    run_id = result.get("run_id")
    if run_id:
        try:
            await save_run({
                "run_id": run_id,
                "conversation_id": conversation_id,
                "question": request.question,
                "answer": answer,
                "citations": [c.model_dump() for c in citations],
                "source_chunks": [sc.model_dump() for sc in source_chunks],
                "verification": verification_out.model_dump(),
                "output_format": result.get("output_format", "CREAC"),
                "reasoning_steps": result.get("reasoning_steps", []),
                "intent": result.get("intent", ""),
                "sources_used": response.sources_used,
                "model_provider": result.get("model_provider", ""),
                "model_name": result.get("model_name", ""),
                "created_at": datetime.utcnow().isoformat(),
                "query_analysis": result.get("query_analysis"),
            })
        except Exception as exc:  # noqa: BLE001
            logger.warning("run persistence failed: %s", exc)

    return response


# ---------------------------------------------------------------------------
# POST /chat — Legacy non-streaming endpoint (backward compatible)
# ---------------------------------------------------------------------------

@router.post("/chat", response_model=ChatResponse)
async def chat_endpoint(request: ChatRequest, http_request: Request):
    """Legacy endpoint — uses original workflow for backward compatibility."""
    await _enforce_rate_limit(http_request)
    _validate_mode(request)
    
    conversation_id = request.conversation_id or str(uuid.uuid4())
    history = await load_history(conversation_id)

    graph = get_workflow()
    state: JuryAIState = {
        "question": request.question,
        "history": history,
        "intent": "",
        "legal_chunks": [],
        "web_results": [],
        "citations": [],
        "answer": "",
        "error": None,
        "use_web_search": request.use_web_search,
        "mode": request.mode,
        "session_id": request.session_id,
        "source_filter": request.sources,
        "as_of_date": request.as_of_date,
    }

    result = await graph.ainvoke(state)
    
    answer = result.get("answer", "")
    verification = dict(result.get("verification") or {})
    citations_raw = result.get("citations", [])
    merged = result.get("merged_evidence") or []
    evidence = list(merged) if merged else list(result.get("legal_chunks") or [])
    
    if settings.VERIFIER_GATE_ENABLED and evidence:
        gated = await gate_answer(
            request.question, answer, evidence, _build_evidence_text(evidence), verification
        )
        if gated.answer != answer:
            answer = gated.answer
            citations_raw = derive_citations(gated.verification, evidence, answer)
        verification = dict(gated.verification)
        verification["blocked"] = gated.blocked
        verification["regenerated"] = gated.regenerated
    
    citations = [CitationOut(**c) for c in citations_raw]
    verified_hashes = verified_content_hashes(citations_raw)
    source_chunks = _source_chunks_from_evidence(evidence, verified_hashes, cited_indices(citations_raw), citations_raw)
    verification_out = VerificationOut(**verification)
    
    await append_turn(conversation_id, request.question, answer)
    schedule_record_turn(
        build_turn_record(
            conversation_id=conversation_id,
            question=request.question,
            answer=answer,
            citations=citations_raw,
            verification=verification,
            evidence=evidence,
            model_provider=result.get("model_provider", ""),
            model_name=result.get("model_name", ""),
            client_id=http_request.client.host if http_request.client else "unknown",
            streaming=False,
            cache_hit=False,
        )
    )
    
    return ChatResponse(
        answer=answer,
        citations=citations,
        source_chunks=source_chunks,
        verification=verification_out,
        intent=result.get("intent", ""),
        sources_used=len(result.get("sources_used", {})) if isinstance(result.get("sources_used"), dict) else (result.get("sources_used") or 0),
        conversation_id=conversation_id,
    )


# ---------------------------------------------------------------------------
# POST /chat/enhanced — Enhanced endpoint with run_id, progress, structured output
# ---------------------------------------------------------------------------

@router.post("/chat/enhanced", response_model=ChatResponse)
async def chat_enhanced_endpoint(request: ChatRequest, http_request: Request):
    """Enhanced endpoint with run_id, structured output (CREAC/IRAC/Brief), query analysis."""
    return await _run_enhanced_chat(request, http_request)


# ---------------------------------------------------------------------------
# SSE streaming endpoint
# ---------------------------------------------------------------------------


async def _stream_generator(request: ChatRequest, client_id: str) -> AsyncGenerator[str, None]:
    """Drive the streaming LangGraph workflow and emit SSE events.

    Order: reasoning_step* → source_chunk* → answer_token* → verification →
    (gate, only if the answer was regenerated/blocked) → done.

    The ``source_chunk`` events fire before the answer exists, so their
    ``verified`` flag is necessarily provisional (always False — grounding
    can't be known until the claim-level verifier has run against the full
    answer). ``done`` carries the corrected, post-verification ``source_chunks``
    and ``citations`` — the frontend already prefers ``done.source_chunks``
    over the accumulated per-event chunks (see ChatPanel.jsx), so this is the
    authoritative signal without any SSE contract change.
    """
    conversation_id = request.conversation_id or str(uuid.uuid4())
    start_time = time.monotonic()

    # Step 1: emit intent classification immediately (before graph runs)
    yield _sse("reasoning_step", {
        "step": "intent_classification",
        "detail": "Classifying query intent and selecting retrieval strategy",
    })

    # Answer cache: replay a cached answer as SSE without re-running the pipeline.
    # Skipped for interact mode — see _validate_mode's docstring: the cache key
    # doesn't carry session_id, so sharing it would leak across sessions.
    cached = (
        await answer_cache_get(request.question, request.use_web_search, scope="stream", sources=request.sources)
        if request.mode != "interact" else None
    )
    if cached is not None:
        cached_chunks = cached.get("source_chunks", [])
        for sc in cached_chunks:
            yield _sse("source_chunk", sc)
        answer_text = cached.get("answer", "")
        words = answer_text.split(" ") if answer_text else []
        for i, word in enumerate(words):
            token = word if i == len(words) - 1 else word + " "
            yield _sse("answer_token", {"token": token})
        verification = cached.get("verification") or {}
        cached_citations = cached.get("citations", [])
        yield _sse("verification", verification)
        yield _sse("done", {
            "conversation_id": conversation_id,
            "intent": cached.get("intent", "legal"),
            "sources_used": len(cached_chunks),
            "verification": verification,
            "source_chunks": cached_chunks,
            "citations": cached_citations,
            "cached": True,
        })
        schedule_record_turn(
            build_turn_record(
                conversation_id=conversation_id,
                question=request.question,
                answer=answer_text,
                citations=cached_citations,
                verification=verification,
                evidence=cached_citations,
                model_provider="",
                model_name="",
                client_id=client_id,
                streaming=True,
                cache_hit=True,
            )
        )
        return

    history = await load_history(request.conversation_id)
    state: JuryAIState = {
        "question": request.question,
        "history": history,
        "intent": "",
        "legal_chunks": [],
        "web_results": [],
        "citations": [],
        "answer": "",
        "error": None,
        "use_web_search": request.use_web_search,
        "reasoning_steps": [],
        "web_evidence": [],
        "merged_evidence": [],
        # Streaming mode: generate_answer_node stores the prompt but skips LLM;
        # the SSE generator calls get_llm().astream() for real token streaming.
        "streaming": True,
        "mode": request.mode,
        "session_id": request.session_id,
        "output_format": request.output_format,
        "source_filter": request.sources,
        "as_of_date": request.as_of_date,
    }

    try:
        graph = get_streaming_workflow()

        # Not astream_events: langgraph is unpinned (>=0.2.0) so its event
        # schema isn't guaranteed stable, and route_and_retrieve fans out to
        # legal_retrieve_node/web_search_node via asyncio.gather under one
        # outer "retrieve" node — node-level events couldn't see those inner
        # sub-steps anyway. A plain callback threaded through state (queued
        # here, drained below) reports the moment each step actually happens
        # regardless of graph structure or LangGraph version.
        queue: asyncio.Queue = asyncio.Queue()
        state["on_step"] = queue.put_nowait
        task = asyncio.create_task(graph.ainvoke(state))

        while not task.done():
            try:
                step = await asyncio.wait_for(queue.get(), timeout=0.1)
            except asyncio.TimeoutError:
                continue
            yield _sse("reasoning_step", {**step})

        result = await task
        while not queue.empty():
            step = queue.get_nowait()
            yield _sse("reasoning_step", {**step})

        # Emit source chunks — prefer merged_evidence (full mixed-domain
        # evidence), else fall back to legal_chunks for backward compat with
        # mocked tests. verified is always False here (provisional — see
        # docstring above).
        merged = result.get("merged_evidence", [])
        if merged:
            verify_evidence = merged
            source_chunks = _source_chunks_from_evidence(verify_evidence, set())
        else:
            verify_evidence = result.get("legal_chunks", [])
            source_chunks, _ = _build_source_chunks(result)

        for chunk in source_chunks:
            yield _sse("source_chunk", chunk.model_dump())

        # Stream answer tokens (real token streaming when a prompt is available).
        answer_prompt: str = result.get("answer_prompt", "")
        full_answer = ""
        model_provider = ""
        model_name = ""
        usage: dict = {}
        if answer_prompt:
            try:
                async for llm_chunk in get_llm().astream([HumanMessage(content=answer_prompt)]):
                    token: str = getattr(llm_chunk, "content", "") or ""
                    if token:
                        full_answer += token
                        yield _sse("answer_token", {"token": token})
                    chunk_meta = getattr(llm_chunk, "response_metadata", None) or {}
                    if chunk_meta.get("model_provider"):
                        model_provider = chunk_meta["model_provider"]
                        model_name = chunk_meta.get("model_name", "")
                    if chunk_meta.get("usage"):
                        usage.update(chunk_meta["usage"])
            except Exception as exc:
                # Real gateway failures mid-stream were previously silent — the
                # verifier gate below catches the resulting low-groundedness
                # answer and regenerates, but without this the root cause of a
                # recurring gateway issue would be undiagnosable.
                logger.warning(
                    "LLM stream failed for conversation_id=%s: %s",
                    conversation_id, exc, exc_info=True,
                )
                yield _sse("reasoning_step", {
                    "step": "stream_interrupted",
                    "detail": "Answer generation was interrupted; the verifier gate below will regenerate or block this answer.",
                })
        else:
            answer: str = result.get("answer", "")
            full_answer = answer
            words = answer.split(" ") if answer else []
            for i, word in enumerate(words):
                token = word if i == len(words) - 1 else word + " "
                yield _sse("answer_token", {"token": token})

        # Meta-verification against the full merged evidence (verify_evidence
        # was already resolved above, when the provisional source_chunks were
        # built); _build_evidence_text still prioritizes internal-domain items.
        yield _sse("reasoning_step", {
            "step": "verifying_answer",
            "detail": "Verifying answer groundedness against retrieved evidence",
        })
        verification = await verify_answer(full_answer, verify_evidence)

        # Verifier gate: regenerate / block if the streamed answer is ungrounded.
        # Common case (answer passes) is a no-op, so streaming stays fast.
        if settings.VERIFIER_GATE_ENABLED and verify_evidence:
            gated = await gate_answer(
                request.question, full_answer, verify_evidence,
                _build_evidence_text(verify_evidence), verification,
            )
            verification = dict(gated.verification)
            verification["blocked"] = gated.blocked
            verification["regenerated"] = gated.regenerated
            if gated.answer != full_answer:
                # Answer changed after release — tell the client to replace it.
                full_answer = gated.answer
                if gated.model_provider:
                    model_provider, model_name = gated.model_provider, gated.model_name
                yield _sse("gate", {
                    "answer": full_answer,
                    "blocked": gated.blocked,
                    "regenerated": gated.regenerated,
                })

        yield _sse("verification", verification)

        # Final, post-verification citations/sources — the unified signal,
        # now that the claim-level verdict for the (possibly gated) answer exists.
        final_citations = derive_citations(verification, verify_evidence, full_answer)
        verified_hashes = verified_content_hashes(final_citations)
        final_source_chunks = _source_chunks_from_evidence(
            verify_evidence, verified_hashes, cited_indices(final_citations)
        )

        yield _sse("done", {
            "conversation_id": conversation_id,
            "intent": result.get("intent", "legal"),
            "answer": full_answer,
            "sources_used": len(final_source_chunks),
            "verification": verification,
            "source_chunks": [sc.model_dump() for sc in final_source_chunks],
            "citations": [CitationOut(**c).model_dump() for c in final_citations],
            "elapsed_seconds": round(time.monotonic() - start_time, 1),
            "usage": usage,
        })

        # Audit every resolved turn — including blocked/empty ones, since those
        # are exactly what a compliance review needs a record of. Unlike the
        # cache/memory writes below, this is never conditioned on the answer
        # having shipped successfully.
        schedule_record_turn(
            build_turn_record(
                conversation_id=conversation_id,
                question=request.question,
                answer=full_answer,
                citations=final_citations,
                verification=verification,
                evidence=verify_evidence,
                model_provider=model_provider,
                model_name=model_name,
                client_id=client_id,
                streaming=True,
            )
        )

        # Persist + cache the FINAL (post-gate) answer. Skip empties/blocked.
        if full_answer.strip() and not verification.get("blocked"):
            await append_turn(request.conversation_id, request.question, full_answer)
        if full_answer.strip() and not verification.get("blocked") and request.mode != "interact":
            await answer_cache_set(
                request.question,
                request.use_web_search,
                {
                    "answer": full_answer,
                    "source_chunks": [sc.model_dump() for sc in final_source_chunks],
                    "citations": [CitationOut(**c).model_dump() for c in final_citations],
                    "verification": verification,
                    "intent": result.get("intent", "legal"),
                },
                scope="stream",
                sources=request.sources,
            )
    except Exception as e:
        logger.exception("chat stream failed")
        yield _sse("done", {
            "conversation_id": conversation_id,
            "intent": "error",
            "sources_used": 0,
            "error": str(e),
        })


@router.post("/chat/stream")
async def chat_stream_endpoint(request: ChatRequest, http_request: Request) -> StreamingResponse:
    """SSE endpoint — streams reasoning steps, source chunks, answer tokens, and done."""
    await _enforce_rate_limit(http_request)
    _validate_mode(request)
    return StreamingResponse(
        content=_stream_generator(request, _client_id(http_request)),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )
@router.post("/chat/devils-advocate", response_model=DevilsAdvocateResponse)
async def devils_advocate_endpoint(request: DevilsAdvocateRequest, http_request: Request):
    """Generate the strongest counterargument against a previously saved run's answer."""
    await _enforce_rate_limit(http_request)
    from app.core.graph.nodes import generate_devils_advocate

    run = await load_run(request.run_id)
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")

    counterargument = await generate_devils_advocate(
        run.get("question", ""), run.get("answer", ""), run.get("source_chunks", [])
    )
    return DevilsAdvocateResponse(counterargument=counterargument)


@router.post("/chat/analyze", response_model=ChatAnalyzeResponse)
async def chat_analyze_endpoint(request: ChatRequest, http_request: Request):
    """Analyze query quality and suggest improvements (v1 parity)."""
    await _enforce_rate_limit(http_request)
    _validate_mode(request)
    
    from app.core.graph.nodes import query_analysis_node
    from app.core.graph.state import JuryAIState
    
    state: JuryAIState = {
        "question": request.question,
        "history": [],
        "intent": "",
        "legal_chunks": [],
        "web_results": [],
        "citations": [],
        "answer": "",
        "error": None,
        "use_web_search": False,
        "mode": "ask",
        "session_id": None,
    }
    
    result = await query_analysis_node(state)
    analysis = result.get("query_analysis", {})
    
    return ChatAnalyzeResponse(
        original_question=request.question,
        analysis=QueryAnalysisOut(**analysis),
        should_improve=analysis.get("score", 5) < 7,
    )


# ---------------------------------------------------------------------------
# POST /chat/improve — Get suggested rewrite for a query
# ---------------------------------------------------------------------------

@router.post("/chat/improve", response_model=ChatImproveResponse)
async def chat_improve_endpoint(request: ChatRequest, http_request: Request):
    """Get suggested rewrite for a query without running full research."""
    await _enforce_rate_limit(http_request)
    _validate_mode(request)
    
    from app.core.graph.nodes import query_analysis_node, prompt_improvement_node
    from app.core.graph.state import JuryAIState
    
    state: JuryAIState = {
        "question": request.question,
        "history": [],
        "intent": "",
        "legal_chunks": [],
        "web_results": [],
        "citations": [],
        "answer": "",
        "error": None,
        "use_web_search": False,
        "mode": "ask",
        "session_id": None,
    }
    
    analysis_result = await query_analysis_node(state)
    state.update(analysis_result)
    improve_result = await prompt_improvement_node(state)
    
    improved_q = improve_result.get("question", request.question)
    
    return ChatImproveResponse(
        original_question=request.question,
        improved_question=improved_q,
        analysis=QueryAnalysisOut(**analysis_result.get("query_analysis", {})),
        was_improved=improved_q != request.question,
    )


# ---------------------------------------------------------------------------
# POST /chat/download — PDF download (with/without citations)
# ---------------------------------------------------------------------------

@router.post("/chat/download")
async def download_answer(request: ChatRequest, http_request: Request):
    """Generate and return a ZIP with two PDFs: with and without citations."""
    import zipfile
    
    await _enforce_rate_limit(http_request)
    _validate_mode(request)
    conversation_id = request.conversation_id or str(uuid.uuid4())

    # Run enhanced chat to get full result
    result = await _run_enhanced_chat(request, http_request)
    
    answer = result.answer
    verification = result.verification
    citations = result.citations
    source_chunks = result.source_chunks
    
    pdf_with = _generate_pdf(request.question, answer, citations, source_chunks, verification, include_citations=True)
    pdf_without = _generate_pdf(request.question, answer, citations, source_chunks, verification, include_citations=False)

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("answer_with_citations.pdf", pdf_with)
        zf.writestr("answer_without_citations.pdf", pdf_without)

    zip_buffer.seek(0)
    return Response(
        content=zip_buffer.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="juryai_answer_{conversation_id[:8]}.zip"'},
    )


@router.post("/chat/download-pdf")
async def download_answer_pdf(request: ChatDownloadPdfRequest):
    """Generate a single PDF from already-generated answer data.

    Unlike /chat/download, this does NOT re-run the chat pipeline. It accepts
    the answer, citations, source_chunks, and verification directly and
    produces a PDF — fast and reliable.
    """
    verification = VerificationOut(**request.verification) if request.verification else None
    pdf_bytes = _generate_pdf(
        question=request.question,
        answer=request.answer,
        citations=[CitationOut(**c) for c in request.citations],
        source_chunks=[SourceChunkOut(**sc) for sc in request.source_chunks],
        verification=verification,
        include_citations=request.include_citations,
    )
    suffix = "with-citations" if request.include_citations else "plain"
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="juryai-answer-{suffix}.pdf"'},
    )


@router.post("/chat/download-docx")
async def download_answer_docx(request: ChatDownloadPdfRequest):
    """Generate a single DOCX from already-generated answer data.

    Mirrors /chat/download-pdf — same request shape, does not re-run the
    chat pipeline.
    """
    verification = VerificationOut(**request.verification) if request.verification else None
    docx_bytes = _generate_docx(
        question=request.question,
        answer=request.answer,
        citations=[CitationOut(**c) for c in request.citations],
        source_chunks=[SourceChunkOut(**sc) for sc in request.source_chunks],
        verification=verification,
        include_citations=request.include_citations,
    )
    suffix = "with-citations" if request.include_citations else "plain"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="juryai-answer-{suffix}.docx"'},
    )


# ---------------------------------------------------------------------------
# Helpers for streaming (moved here to avoid circular imports)
# ---------------------------------------------------------------------------

async def route_and_retrieve_with_progress(state: JuryAIState) -> dict:
    from app.core.graph.workflow import emit_progress
    emit_progress(state, "internal_retrieval")
    from app.core.graph.workflow import route_and_retrieve
    result = await route_and_retrieve(state)
    emit_progress(state, "web_search" if state.get("use_web_search") else "evidence_merge")
    return result


async def _enforce_rate_limit(request: Request):
    client_ip = request.client.host if request.client else "unknown"
    await check_rate_limit(client_ip)


def _validate_mode(request: ChatRequest):
    if request.mode not in ("ask", "interact"):
        raise HTTPException(status_code=400, detail="mode must be 'ask' or 'interact'")


# ---------------------------------------------------------------------------
# Run storage routes
# ---------------------------------------------------------------------------

class SourceActionRequest(BaseModel):
    action: str = "view"  # "view" | "download" | "copy_citation" | "read_chunk" | "copy_chunk" | "open_window"


class FollowUpRequest(BaseModel):
    question: str
    use_web_search: bool = False


@router.get("/runs/{run_id}")
async def get_run(run_id: str):
    """Return a persisted run payload."""
    data = await load_run(run_id)
    if not data:
        raise HTTPException(status_code=404, detail="Run not found")
    return data


@router.post("/runs/{run_id}/followup")
async def follow_up_run(run_id: str, request: FollowUpRequest, http_request: Request):
    """Run a follow-up question linked to an existing run's conversation."""
    parent = await load_run(run_id)
    if not parent:
        raise HTTPException(status_code=404, detail="Parent run not found")

    chat_request = ChatRequest(
        question=request.question,
        conversation_id=parent.get("conversation_id"),
        use_web_search=request.use_web_search,
        mode="ask",
        output_format=parent.get("output_format", "CREAC"),
    )
    response = await _run_enhanced_chat(chat_request, http_request)
    return response


@router.post("/runs/{run_id}/sources/{source_index}/actions")
async def source_action(run_id: str, source_index: int, request: SourceActionRequest):
    """Perform an action on a source chunk."""
    data = await load_run(run_id)
    if not data:
        raise HTTPException(status_code=404, detail="Run not found")

    chunks = data.get("source_chunks", [])
    if source_index < 0 or source_index >= len(chunks):
        raise HTTPException(status_code=404, detail="Source not found")

    chunk = chunks[source_index]
    action = request.action

    if action == "view":
        return {"chunk": chunk}

    if action == "copy_citation":
        quote = chunk.get("citation_quote") or chunk.get("text", "")
        source = chunk.get("source", "")
        page = chunk.get("page")
        citation = f'"{quote}"' if quote else ""
        ref = f"{source}" + (f" p.{page + 1}" if page is not None else "")
        return {"citation": f"{citation} — {ref}".strip(" —")}

    if action == "download":
        text = chunk.get("text", "")
        source = chunk.get("source", "source")
        filename = f"{source.replace('/', '_')}.txt"
        return Response(
            content=text.encode("utf-8"),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if action == "read_chunk":
        return {"text": chunk.get("text", "")}

    if action == "copy_chunk":
        return {"text": chunk.get("text", "")}

    if action == "open_window":
        url = chunk.get("url") or chunk.get("source", "")
        return {"url": url}

    raise HTTPException(status_code=400, detail=f"Unsupported action: {action}")


@router.get("/runs/{run_id}/source/{source_index}")
async def get_run_source(run_id: str, source_index: int):
    """Return a single source chunk for a run."""
    data = await load_run(run_id)
    if not data:
        raise HTTPException(status_code=404, detail="Run not found")

    chunks = data.get("source_chunks", [])
    if source_index < 0 or source_index >= len(chunks):
        raise HTTPException(status_code=404, detail="Source not found")

    return chunks[source_index]


@router.get("/runs/{run_id}/source/{source_index}/file")
async def get_run_source_file(run_id: str, source_index: int):
    """Download the source document for a run chunk."""
    data = await load_run(run_id)
    if not data:
        raise HTTPException(status_code=404, detail="Run not found")

    chunks = data.get("source_chunks", [])
    if source_index < 0 or source_index >= len(chunks):
        raise HTTPException(status_code=404, detail="Source not found")

    chunk = chunks[source_index]
    text = chunk.get("text", "")
    source = chunk.get("source", "source")
    filename = f"{source.replace('/', '_')}.txt"
    return Response(
        content=text.encode("utf-8"),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/runs/{run_id}/download")
async def download_run(run_id: str):
    """Generate and return a PDF for a persisted run."""
    data = await load_run(run_id)
    if not data:
        raise HTTPException(status_code=404, detail="Run not found")

    verification_raw = data.get("verification") or {}
    verification = VerificationOut(**verification_raw) if verification_raw else None
    pdf_bytes = _generate_pdf(
        question=data.get("question", ""),
        answer=data.get("answer", ""),
        citations=data.get("citations", []),
        source_chunks=data.get("source_chunks", []),
        verification=verification,
        include_citations=True,
    )
    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="run-{run_id}.pdf"'},
    )
