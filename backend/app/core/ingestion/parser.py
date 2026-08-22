from __future__ import annotations
import logging
from typing import TypedDict
import pymupdf
import pymupdf4llm

logger = logging.getLogger(__name__)

CHUNK_SIZE = 512
CHUNK_OVERLAP = 64
_HEADING_PREFIX = "#"
_TEXT_SUFFIXES = (".txt", ".md")
_EXCEL_SUFFIXES = (".xlsx", ".xlsm")
_CSV_SUFFIXES = (".csv",)
# .docx only — legacy binary .doc requires a different parser entirely
# (python-docx can't read it) and isn't supported.
_DOCX_SUFFIXES = (".docx",)
_IMAGE_SUFFIXES = (".jpg", ".jpeg", ".png", ".webp", ".tiff", ".tif", ".bmp")


class Chunk(TypedDict):
    text: str
    page: int
    source: str


def _new(text: str, source: str, page: int) -> Chunk:
    return {"text": text, "page": page, "source": source}


def _char_split(text: str, source: str, page: int) -> list[Chunk]:
    """Overflow fallback: slice an oversized section into overlapping windows."""
    chunks, start = [], 0
    while start < len(text):
        piece = text[start:start + CHUNK_SIZE].strip()
        if piece:
            chunks.append(_new(piece, source, page))
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def _split_sections(markdown: str) -> list[str]:
    """Split markdown at heading boundaries so each section keeps its heading."""
    sections: list[str] = []
    current: list[str] = []
    for line in markdown.splitlines(keepends=True):
        if line.lstrip().startswith(_HEADING_PREFIX) and current:
            sections.append("".join(current))
            current = [line]
        else:
            current.append(line)
    if current:
        sections.append("".join(current))
    return sections


def _chunk_markdown(markdown: str, source: str, page: int) -> list[Chunk]:
    """Chunk on markdown structure; small sections stay whole (tables/headings
    survive intact), oversized ones fall back to overlapping character windows."""
    chunks: list[Chunk] = []
    for section in _split_sections(markdown):
        section = section.strip()
        if not section:
            continue
        if len(section) <= CHUNK_SIZE:
            chunks.append(_new(section, source, page))
        else:
            chunks.extend(_char_split(section, source, page))
    return chunks


def _parse_pdf(data: bytes, filename: str) -> list[Chunk]:
    """Parse PDF with pymupdf4llm, fallback to OCR if text extraction is poor."""
    doc = pymupdf.open(stream=data, filetype="pdf")
    chunks: list[Chunk] = []

    # First pass: try standard markdown extraction
    for page_index, page in enumerate(pymupdf4llm.to_markdown(doc, page_chunks=True)):
        page_text = page["text"]
        # Check if page has meaningful text
        if len(page_text.strip()) >= 50:
            chunks.extend(_chunk_markdown(page_text, filename, page_index))
        else:
            # Page appears to be scanned/image-based — try OCR
            ocr_chunks = _parse_page_ocr(data, filename, page_index)
            if ocr_chunks:
                chunks.extend(ocr_chunks)
                logger.info("OCR used for page %d of %s", page_index + 1, filename)
            else:
                # Fallback to whatever text we got
                chunks.extend(_chunk_markdown(page_text, filename, page_index))

    return chunks


def _parse_page_ocr(data: bytes, filename: str, page_index: int) -> list[Chunk]:
    """OCR a single scanned page with pdf2image (render) + pytesseract."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        from app.config.settings import settings
    except ImportError:
        logger.debug("pytesseract/pdf2image not available, skipping OCR")
        return []

    if not settings.OCR_ENABLED:
        return []

    try:
        images = convert_from_bytes(
            data, dpi=settings.OCR_DPI,
            first_page=page_index + 1, last_page=page_index + 1,
        )
        if not images:
            return []
        text = pytesseract.image_to_string(images[0], lang=settings.OCR_LANG)
        if len(text.strip()) < settings.OCR_MIN_CHARS_PER_PAGE:
            return []
        return _chunk_markdown(text, filename, page_index)
    except Exception as exc:
        logger.warning("OCR failed for page %d of %s: %s", page_index + 1, filename, exc)
        return []


def _rows_to_markdown(sheet: str, rows: list[list[str]]) -> str:
    """Render worksheet rows as a markdown table under a sheet-name heading."""
    header, *body = rows
    width = len(header)
    lines = [
        f"# {sheet}",
        "",
        "| " + " | ".join(header) + " |",
        "| " + " | ".join("---" for _ in header) + " |",
    ]
    for row in body:
        cells = (row + [""] * width)[:width]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _parse_excel(data: bytes, filename: str) -> list[Chunk]:
    """Extract each worksheet as a markdown table (one page per sheet)."""
    from io import BytesIO
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    try:
        chunks: list[Chunk] = []
        for page_index, ws in enumerate(wb.worksheets):
            rows = [
                ["" if cell is None else str(cell) for cell in row]
                for row in ws.iter_rows(values_only=True)
            ]
            if not any(any(cell.strip() for cell in row) for row in rows):
                continue
            markdown = _rows_to_markdown(ws.title, rows)
            chunks.extend(_chunk_markdown(markdown, filename, page_index))
        return chunks
    finally:
        wb.close()


def _parse_csv(data: bytes, filename: str) -> list[Chunk]:
    """Render CSV rows as a single markdown table — same shape as an Excel sheet."""
    import csv
    from io import StringIO

    # utf-8-sig strips a BOM if present (common from Excel-exported CSVs)
    # without corrupting files that don't have one.
    text = data.decode("utf-8-sig", errors="replace")
    rows = [row for row in csv.reader(StringIO(text)) if any(cell.strip() for cell in row)]
    if not rows:
        return []
    markdown = _rows_to_markdown(filename.rsplit("/", 1)[-1], rows)
    return _chunk_markdown(markdown, filename, 0)


def _parse_docx(data: bytes, filename: str) -> list[Chunk]:
    """Extract paragraphs (with heading levels preserved) and tables from a
    Word document. Tables are appended after the body text rather than
    interleaved at their original position — a reasonable simplification for
    typical legal documents, where tables are usually schedules/appendices
    rather than content threaded through the prose."""
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        if style.lower().startswith("heading"):
            level = next((c for c in style if c.isdigit()), "1")
            parts.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [row for row in rows if any(cell for cell in row)]
        if rows:
            parts.append(_rows_to_markdown("Table", rows))

    if not parts:
        return []
    return _chunk_markdown("\n\n".join(parts), filename, 0)


def _parse_image(data: bytes, filename: str) -> list[Chunk]:
    """OCR a standalone image (a photographed or scanned page saved directly
    as an image, not embedded in a PDF) — same pytesseract stack _parse_pdf
    already uses per-page, applied directly to the whole image."""
    try:
        import pytesseract
        from io import BytesIO
        from PIL import Image
        from app.config.settings import settings
    except ImportError:
        logger.debug("pytesseract/Pillow not available, skipping image OCR")
        return []

    if not settings.OCR_ENABLED:
        return []

    try:
        image = Image.open(BytesIO(data))
        text = pytesseract.image_to_string(image, lang=settings.OCR_LANG)
        if len(text.strip()) < settings.OCR_MIN_CHARS_PER_PAGE:
            return []
        return _chunk_markdown(text, filename, 0)
    except Exception as exc:
        logger.warning("Image OCR failed for %s: %s", filename, exc)
        return []


def parse_bytes(data: bytes, filename: str) -> list[Chunk]:
    lower = filename.lower()
    if lower.endswith(_CSV_SUFFIXES):
        try:
            return _parse_csv(data, filename)
        except Exception as exc:
            logger.warning("CSV parse failed for %s: %s", filename, exc)
            return []
    if lower.endswith(_DOCX_SUFFIXES):
        try:
            return _parse_docx(data, filename)
        except Exception as exc:
            logger.warning("DOCX parse failed for %s: %s", filename, exc)
            return []
    if lower.endswith(_IMAGE_SUFFIXES):
        return _parse_image(data, filename)
    if lower.endswith(_TEXT_SUFFIXES):
        return _chunk_markdown(data.decode("utf-8", errors="replace"), filename, 0)
    if lower.endswith(_EXCEL_SUFFIXES):
        try:
            return _parse_excel(data, filename)
        except Exception as exc:
            logger.warning("Excel parse failed for %s: %s", filename, exc)
            return []
    try:
        return _parse_pdf(data, filename)
    except (pymupdf.FileDataError, ValueError, RuntimeError) as exc:
        logger.warning("PDF parse failed for %s, falling back to text: %s", filename, exc)
        return _chunk_markdown(data.decode("utf-8", errors="replace"), filename, 0)